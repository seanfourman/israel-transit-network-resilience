"""Build an editable Hebrew DOCX of the final report from final_report_he.html.

The HTML is the single source of truth; this parses it and emits a Word document
with right-to-left paragraphs, the David font, embedded figures and the data
table, so the report can be edited and re-paginated in Word.

Run:  python reports/build_report_docx.py            # -> final_report_he.docx
      python reports/build_report_docx.py final_report_he_extended
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString
from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

REPO = Path(__file__).resolve().parent.parent
STEM = Path(sys.argv[1]).stem if len(sys.argv) > 1 else "final_report_he"
SRC = REPO / "reports" / f"{STEM}.html"
OUT = REPO / "reports" / f"{STEM}.docx"
FIG_ROOT = REPO / "outputs" / "nb"

# Monochrome by default: a formal black-and-white academic look (like a thesis).
# Set MONO = False for the navy-accented styling.
MONO = True
BLACK = RGBColor(0x00, 0x00, 0x00)
NAVY = BLACK if MONO else RGBColor(0x12, 0x31, 0x4F)
GREY = RGBColor(0x33, 0x33, 0x33) if MONO else RGBColor(0x5A, 0x6B, 0x7D)
HEAD_FILL = "D9D9D9" if MONO else "12314F"   # table-header shading
HEAD_TEXT = BLACK if MONO else RGBColor(0xFF, 0xFF, 0xFF)
BODY_FONT = "David"
HEAD_FONT = "David" if MONO else "Arial"
BODY_PT = 12          # recommended body size
LINE = 1.15           # compact-ish so page count is easy to gauge; 1.5 for a formal norm

FIG_TOKEN = re.compile(r"\{\{FIG:([^}]+)\}\}")


def _style_run(run, size=BODY_PT, bold=False, italic=False, color=None, font=BODY_FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), font)
    szcs = OxmlElement("w:szCs")
    szcs.set(qn("w:val"), str(int(size * 2)))
    rpr.append(szcs)
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rpr.append(rtl)


def _rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)


def _bottom_border(paragraph, color="000000" if MONO else "12314F", sz="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_title_page(doc):
    """A conventional Hebrew academic title page (עמוד שער), on its own page."""
    def cpar(text, size, bold=False, before=0, after=6, font=BODY_FONT, color=BLACK):
        p = doc.add_paragraph()
        _rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        _style_run(p.add_run(text), size=size, bold=bold, color=color, font=font)
        return p

    cpar("‹שם המוסד / הפקולטה›", 12, before=70, color=GREY)
    cpar("פרויקט גמר · הקורס: אלגוריתמים בגרפים", 12, after=6, color=GREY)
    cpar("מהי תחנה קריטית?", 30, bold=True, before=95)
    cpar("מסגרת השוואתית לניתוח מרכזיות ועמידות ברשת התחבורה הציבורית בישראל", 14, before=4, after=8)
    hr = doc.add_paragraph()
    hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _bottom_border(hr)
    hr.paragraph_format.space_after = Pt(4)
    cpar("מגישים: ‹שם מלא›, ‹שם מלא›, ‹שם מלא›", 12, before=80)
    cpar("מנחה: ד\"ר אבנר פריאל", 12)
    cpar("מקור הנתונים: פיד ה-GTFS של משרד התחבורה (אפריל 2026)", 11, before=6, color=GREY)
    cpar("יולי 2026", 12, before=26)
    doc.add_page_break()


def _collapse(text: str) -> str:
    return re.sub(r"\s+", " ", text)


def add_inline(paragraph, node, size=BODY_PT, color=None):
    """Render an element's inline children as runs, keeping bold/italic."""
    for child in node.children:
        if isinstance(child, NavigableString):
            t = _collapse(str(child))
            if t:
                _style_run(paragraph.add_run(t), size=size, color=color)
        elif child.name in ("b", "strong"):
            _style_run(paragraph.add_run(_collapse(child.get_text())), size=size, bold=True, color=color)
        elif child.name in ("em", "i"):
            _style_run(paragraph.add_run(_collapse(child.get_text())), size=size, italic=True, color=color)
        elif child.name == "span":
            cls = child.get("class", []) or []
            _style_run(paragraph.add_run(_collapse(child.get_text())), size=size,
                       bold=("fnum" in cls), color=(NAVY if "fnum" in cls else color))
        else:
            _style_run(paragraph.add_run(_collapse(child.get_text())), size=size, color=color)


def body_para(doc, node, size=BODY_PT, color=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
    p = doc.add_paragraph()
    _rtl(p)
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = LINE
    pf.space_after = Pt(space_after)
    add_inline(p, node, size=size, color=color)
    return p


def heading(doc, text, size=13.5, border=True, space_before=12):
    p = doc.add_paragraph()
    _rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(4)
    _style_run(p.add_run(text), size=size, bold=True, color=NAVY, font=HEAD_FONT)
    if border:
        _bottom_border(p)
    return p


def add_figure(doc, figure):
    m = FIG_TOKEN.search(str(figure))
    if not m:
        return
    path = FIG_ROOT / m.group(1).strip()
    if not path.exists():
        return
    w, h = Image.open(path).size
    width_cm = 15.5
    if width_cm * h / w > 16.5:            # keep tall figures on the page
        width_cm = 16.5 * w / h
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.space_before = Pt(6)
    pic_p.paragraph_format.space_after = Pt(2)
    pic_p.add_run().add_picture(str(path), width=Cm(width_cm))
    cap = figure.find("figcaption")
    if cap:
        cp = doc.add_paragraph()
        _rtl(cp)
        cp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        cp.paragraph_format.space_after = Pt(10)
        add_inline(cp, cap, size=9.5, color=RGBColor(0x39, 0x42, 0x4C))


def add_table(doc, table_el):
    cap = table_el.find("caption")
    if cap:
        cp = doc.add_paragraph()
        _rtl(cp)
        cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cp.paragraph_format.space_after = Pt(2)
        _style_run(cp.add_run(_collapse(cap.get_text())), size=10, bold=True, color=NAVY, font="Arial")
    rows = table_el.find_all("tr")
    ncol = max(len(r.find_all(["td", "th"])) for r in rows)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = "Table Grid"
    t.alignment = 2  # right
    # force table RTL
    tblPr = t._tbl.tblPr
    bidi = OxmlElement("w:bidiVisual")
    tblPr.append(bidi)
    for r in rows:
        cells = r.find_all(["td", "th"])
        row = t.add_row().cells
        for i, c in enumerate(cells):
            is_head = c.name == "th"
            cell = row[i]
            cell.paragraphs[0].clear()
            _rtl(cell.paragraphs[0])
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].add_run(_collapse(c.get_text()))
            _style_run(run, size=10, bold=is_head, color=(HEAD_TEXT if is_head else None))
            if is_head:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), HEAD_FILL)
                cell._tc.get_or_add_tcPr().append(shd)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_list(doc, list_el, ordered):
    for i, li in enumerate(list_el.find_all("li", recursive=False), 1):
        p = doc.add_paragraph()
        _rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = LINE
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.6)
        prefix = f"{i}. " if ordered else "• "
        _style_run(p.add_run(prefix), bold=ordered)
        add_inline(p, li)


def main():
    soup = BeautifulSoup(SRC.read_text(encoding="utf-8"), "lxml")
    doc = Document()

    # page + default style
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    sec.left_margin = sec.right_margin = Cm(2.0)
    sec._sectPr.append(OxmlElement("w:bidi"))  # RTL section
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_PT)

    # A formal document opens with a dedicated title page.
    add_title_page(doc)

    abs = soup.find("div", class_="abstract")
    if abs:
        for p in abs.find_all("p"):
            body_para(doc, p, size=11)

    # main flow: headings, paragraphs, figures, tables, lists
    body = soup.find("body")
    for el in body.find_all(["h2", "p", "figure", "table", "ol", "ul"], recursive=True):
        # skip elements already consumed inside head/abstract
        if el.find_parent("div", class_=["head", "abstract"]):
            continue
        if el.name == "h2":
            heading(doc, _collapse(el.get_text()))
        elif el.name == "p":
            small = "small" in (el.get("class", []) or [])
            body_para(doc, el, size=9.5 if small else BODY_PT,
                      color=GREY if small else None)
        elif el.name == "figure":
            add_figure(doc, el)
        elif el.name == "table":
            add_table(doc, el)
        elif el.name in ("ol", "ul"):
            add_list(doc, el, ordered=(el.name == "ol"))

    doc.save(OUT)
    print(f"wrote {OUT.name}  ({OUT.stat().st_size/1024:.0f} KB)")


if __name__ == "__main__":
    main()
